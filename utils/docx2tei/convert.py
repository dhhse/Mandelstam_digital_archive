from docx import Document
from utils.teiwriter.teiwriter import Tei
import re
import os


def get_contains_words():
    regex = re.compile("[a-zA-ZА-Яа-я]")
    return lambda s: regex.match(s)


contains_words = get_contains_words()


def text2tei_body(text):
    lines = text.splitlines()
    # Remove lines without words
    lines = [line for line in lines if contains_words(line)]

    quatrains = []
    line_n = 0
    for quatrain_n in range(int(len(lines) / 4)):
        quatrain = {"type": "quatrain", "lines": []}
        for line in lines[(quatrain_n * 4):((quatrain_n * 4) + 4)]:
            line_n = line_n + 1
            quatrain["lines"].append({
                "n": line_n,
                "text": line
            })
        quatrains.append(quatrain)

    return quatrains, lines[0]

def main():
    for f_name in os.listdir("files/docs"):
        if f_name.endswith(".docx"):
            print(f"open file {f_name}")
            doc = Document(f"files/docs/{f_name}")
            u = 0
            for idx, cell in enumerate(doc.tables[0].rows[1].cells[1:]):
                if len(cell.text.strip()) == 0:
                    continue
                source = doc.tables[0].rows[0].cells[idx + 1].text
                comment = doc.tables[0].rows[2].cells[idx + 1].text
                text, first_line = text2tei_body(cell.text)
                u += 1


                t = Tei()
                t.set_resp_statements([
                    {
                        "resp": 'подготовка TEI/XML',
                        "name": 'Алексей Литвинцев, Вероника Файнберг, Тимофей Молчанов'
                    }, {
                        "resp": 'Идея, постановка задач, руководство',
                        "name": 'Анастасия Бонч-Осмоловская, Павел Нерлер, Леонид Видгоф,  Дмитрий Зуев, Вероника Файнберг, Павел Литвинов',
                        "orgName": 'Центр Digital Humanities НИУ ВШЭ'
                    }
                ])
                t.set_header_title(first_line, 'О.Э. Мандельштам')
                t.set_header_publication('Москва', 'АРТ-БИЗНЕС-ЦЕНТР МОСКВА', '1993')
                t.set_header_source('О.Э. Мандельштам', source)
                t.set_header_encoding('Издание подготовлено Мандельштамовским обществом')
                t.set_header_profile({
                    "notBefore": '1920',
                    "notAfter": '1920',
                    "text": '1920'
                }, [
                    {
                        "ident": "RU",
                        "text": 'Whole text in Russian.'
                    }
                ], "poem")
                t.set_header_note(comment)
                t.set_head(first_line)
                t.set_body_text(text)
                t.set_header_catref("type", "target")
                t.set_header_catref("type2", "target4")
                t.set_header_catref("type3", "target5")
                t.set_body_text_metadata(volume="2", type="poem", type_n="63", rhyme="aabbccddeeffaaff")

                with open(f'files/tei/{f_name.replace(".docx", "")}_{str(int(idx) + 1)}.tei', 'w') as f:
                    f.write(t.export())
                    f.close()
            break


if __name__ == "__main__":
    main()
